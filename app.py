"""
=============================================================================
NeuroBot API v2.1 - PostgreSQL Edition (CORS Fixed)
=============================================================================
"""

from flask import Flask, request, jsonify
from flask_cors import CORS
import psycopg2
from psycopg2 import pool
import json, os, io, hashlib, re, traceback
from datetime import datetime
import urllib.request
from html.parser import HTMLParser

import torch
import torch.nn as nn
import torch.optim as optim
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# =============================================================================
# FLASK APP + CORS (EXPLICIT)
# =============================================================================
app = Flask(__name__)

# CORS explícito: permite CUALQUIER origen
CORS(app, resources={r"/*": {"origins": "*"}}, supports_credentials=False)

# Headers CORS manuales como fallback
@app.after_request
def after_request(response):
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
    response.headers.add('Access-Control-Allow-Methods', 'GET,POST,PUT,DELETE,OPTIONS')
    return response

# =============================================================================
# CONFIGURACIÓN
# =============================================================================
DATABASE_URL = os.environ.get("DATABASE_URL", "")
EMBEDDING_DIM = 256
LEARNING_RATE = 0.001
TOP_K_RESULTS = 5
MIN_CHUNK_LENGTH = 50
MAX_CHUNK_LENGTH = 1000

# =============================================================================
# POOL DE CONEXIONES PostgreSQL
# =============================================================================
db_pool = None

def get_db_pool():
    global db_pool
    if db_pool is None:
        if not DATABASE_URL:
            raise Exception("DATABASE_URL no configurada. Agrega la variable de entorno en Render.")
        dsn = DATABASE_URL
        # Render a veces da postgres:// en vez de postgresql://
        if dsn.startswith("postgres://"):
            dsn = dsn.replace("postgres://", "postgresql://", 1)
        db_pool = pool.ThreadedConnectionPool(minconn=1, maxconn=5, dsn=dsn)
    return db_pool

def get_conn():
    return get_db_pool().getconn()

def put_conn(conn):
    try:
        get_db_pool().putconn(conn)
    except Exception:
        pass


# =============================================================================
# REDES NEURONALES
# =============================================================================
class KnowledgeEncoder(nn.Module):
    def __init__(self, input_dim, embedding_dim=EMBEDDING_DIM):
        super().__init__()
        self.encoder = nn.Sequential(
            nn.Linear(input_dim, 512), nn.ReLU(), nn.Dropout(0.2),
            nn.Linear(512, embedding_dim), nn.Tanh()
        )
        self.decoder = nn.Sequential(
            nn.Linear(embedding_dim, 512), nn.ReLU(), nn.Dropout(0.2),
            nn.Linear(512, input_dim), nn.Sigmoid()
        )

    def encode(self, x):
        return self.encoder(x)

    def forward(self, x):
        z = self.encode(x)
        return self.decoder(z), z


class AttentionRanker(nn.Module):
    def __init__(self, embedding_dim=EMBEDDING_DIM):
        super().__init__()
        self.query_proj = nn.Linear(embedding_dim, embedding_dim)
        self.key_proj = nn.Linear(embedding_dim, embedding_dim)
        self.value_proj = nn.Linear(embedding_dim, embedding_dim)
        self.score_head = nn.Sequential(
            nn.Linear(embedding_dim, 64), nn.ReLU(), nn.Linear(64, 1), nn.Sigmoid()
        )

    def forward(self, query_emb, knowledge_embs):
        q = self.query_proj(query_emb)
        k = self.key_proj(knowledge_embs)
        v = self.value_proj(knowledge_embs)
        scale = q.size(-1) ** 0.5
        attn = torch.matmul(q.unsqueeze(0), k.T) / scale
        attn_weights = torch.softmax(attn, dim=-1)
        context = torch.matmul(attn_weights, v)
        scores = self.score_head(v).squeeze(-1)
        return scores, context, attn_weights


# =============================================================================
# PROCESADORES DE DOCUMENTOS
# =============================================================================
class HTMLTextExtractor(HTMLParser):
    def __init__(self):
        super().__init__()
        self.result = []
        self.skip_tags = {'script', 'style', 'nav', 'footer', 'header'}
        self._skip = False

    def handle_starttag(self, tag, attrs):
        if tag in self.skip_tags:
            self._skip = True

    def handle_endtag(self, tag):
        if tag in self.skip_tags:
            self._skip = False

    def handle_data(self, data):
        if not self._skip and data.strip():
            self.result.append(data.strip())

    def get_text(self):
        return ' '.join(self.result)


def extract_text_from_pdf(file_path):
    try:
        import fitz
        doc = fitz.open(file_path)
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
        return text
    except ImportError:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            return "\n".join(p.extract_text() or "" for p in pdf.pages)


def extract_text_from_docx(file_path):
    from docx import Document
    doc = Document(file_path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            text += "\n" + " | ".join(cell.text for cell in row.cells)
    return text


def extract_text_from_url(url):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0 (NeuroBot/2.0)'})
    with urllib.request.urlopen(req, timeout=15) as resp:
        html = resp.read().decode('utf-8', errors='ignore')
    ext = HTMLTextExtractor()
    ext.feed(html)
    return ext.get_text()


# =============================================================================
# MOTOR DE CONOCIMIENTO
# =============================================================================
class KnowledgeEngine:
    def __init__(self):
        self.vectorizer = TfidfVectorizer(max_features=5000, ngram_range=(1, 3), sublinear_tf=True)
        self.encoder = None
        self.ranker = None
        self.optimizer_enc = None
        self.optimizer_rank = None
        self.is_fitted = False
        self.knowledge_embeddings = None
        self.knowledge_texts = []
        self.knowledge_ids = []
        self.db_ready = False

    def init_db(self):
        """Inicializa BD. Se llama después de que la app arranca."""
        if self.db_ready:
            return True
        try:
            conn = get_conn()
            with conn.cursor() as c:
                c.execute("""CREATE TABLE IF NOT EXISTS knowledge (
                    id VARCHAR(16) PRIMARY KEY, content TEXT NOT NULL,
                    source VARCHAR(500) DEFAULT 'unknown', source_type VARCHAR(50) DEFAULT 'text',
                    chunk_index INTEGER DEFAULT 0, created_at TIMESTAMP DEFAULT NOW(),
                    access_count INTEGER DEFAULT 0, usefulness_score REAL DEFAULT 0.5,
                    metadata JSONB DEFAULT '{}'::jsonb
                )""")
                c.execute("CREATE INDEX IF NOT EXISTS idx_k_type ON knowledge(source_type)")
                c.execute("CREATE INDEX IF NOT EXISTS idx_k_date ON knowledge(created_at DESC)")
                c.execute("""CREATE TABLE IF NOT EXISTS interactions (
                    id SERIAL PRIMARY KEY, question TEXT NOT NULL,
                    answer_chunks JSONB NOT NULL, feedback REAL DEFAULT 0.0,
                    created_at TIMESTAMP DEFAULT NOW()
                )""")
                c.execute("""CREATE TABLE IF NOT EXISTS training_log (
                    id SERIAL PRIMARY KEY, epoch INTEGER, loss REAL,
                    num_samples INTEGER, created_at TIMESTAMP DEFAULT NOW()
                )""")
                c.execute("""CREATE TABLE IF NOT EXISTS model_weights (
                    id VARCHAR(50) PRIMARY KEY, weights BYTEA NOT NULL,
                    updated_at TIMESTAMP DEFAULT NOW()
                )""")
            conn.commit()
            put_conn(conn)
            self.db_ready = True
            print("✅ Tablas PostgreSQL listas")
            self._load_knowledge()
            return True
        except Exception as e:
            print(f"⚠️  Error BD: {e}")
            traceback.print_exc()
            try:
                put_conn(conn)
            except:
                pass
            return False

    def _ensure_db(self):
        """Asegura que la BD esté lista antes de cualquier operación."""
        if not self.db_ready:
            self.init_db()
        return self.db_ready

    def _load_knowledge(self):
        conn = get_conn()
        try:
            with conn.cursor() as c:
                c.execute("SELECT id, content FROM knowledge ORDER BY created_at")
                rows = c.fetchall()
            if rows:
                self.knowledge_ids = [r[0] for r in rows]
                self.knowledge_texts = [r[1] for r in rows]
                self._fit_models()
                print(f"✅ {len(rows)} chunks cargados")
            else:
                print("ℹ️  BD vacía, lista para aprender")
        finally:
            put_conn(conn)

    def _fit_models(self):
        if len(self.knowledge_texts) < 2:
            return
        tfidf_matrix = self.vectorizer.fit_transform(self.knowledge_texts)
        input_dim = tfidf_matrix.shape[1]
        self.encoder = KnowledgeEncoder(input_dim, EMBEDDING_DIM)
        self.ranker = AttentionRanker(EMBEDDING_DIM)
        self.optimizer_enc = optim.Adam(self.encoder.parameters(), lr=LEARNING_RATE)
        self.optimizer_rank = optim.Adam(self.ranker.parameters(), lr=LEARNING_RATE)
        X = torch.FloatTensor(tfidf_matrix.toarray())
        self._train_encoder(X, epochs=50)
        self.encoder.eval()
        with torch.no_grad():
            _, self.knowledge_embeddings = self.encoder(X)
        self.is_fitted = True

    def _train_encoder(self, X, epochs=50):
        self.encoder.train()
        criterion = nn.MSELoss()
        fl = 0.0
        for _ in range(epochs):
            self.optimizer_enc.zero_grad()
            reconstructed, embeddings = self.encoder(X)
            loss = criterion(reconstructed, X)
            if embeddings.size(0) > 1:
                sim = torch.mm(embeddings, embeddings.T)
                loss += torch.mean((sim - torch.eye(embeddings.size(0))) ** 2) * 0.01
            loss.backward()
            self.optimizer_enc.step()
            fl = loss.item()
        try:
            conn = get_conn()
            with conn.cursor() as c:
                c.execute("INSERT INTO training_log (epoch,loss,num_samples) VALUES (%s,%s,%s)", (epochs, fl, X.size(0)))
            conn.commit()
            put_conn(conn)
        except:
            pass

    def _train_from_feedback(self, q_emb, relevant_ids, score):
        if self.ranker is None or self.knowledge_embeddings is None:
            return
        self.ranker.train()
        self.optimizer_rank.zero_grad()
        scores, _, _ = self.ranker(q_emb, self.knowledge_embeddings)
        target = torch.zeros(len(self.knowledge_ids))
        for kid in relevant_ids:
            if kid in self.knowledge_ids:
                target[self.knowledge_ids.index(kid)] = max(0, score)
        nn.BCELoss()(scores, target).backward()
        self.optimizer_rank.step()

    def _chunk_text(self, text):
        text = re.sub(r'\s+', ' ', text).strip()
        if len(text) <= MAX_CHUNK_LENGTH:
            return [text] if len(text) >= MIN_CHUNK_LENGTH else []
        paragraphs = re.split(r'\n\s*\n', text)
        chunks, current = [], ""
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            if len(current) + len(para) <= MAX_CHUNK_LENGTH:
                current += (" " if current else "") + para
            else:
                if len(current) >= MIN_CHUNK_LENGTH:
                    chunks.append(current.strip())
                current = para
        if len(current) >= MIN_CHUNK_LENGTH:
            chunks.append(current.strip())
        final = []
        for chunk in chunks:
            if len(chunk) > MAX_CHUNK_LENGTH:
                sentences = re.split(r'(?<=[.!?])\s+', chunk)
                sub = ""
                for s in sentences:
                    if len(sub) + len(s) <= MAX_CHUNK_LENGTH:
                        sub += (" " if sub else "") + s
                    else:
                        if len(sub) >= MIN_CHUNK_LENGTH:
                            final.append(sub.strip())
                        sub = s
                if len(sub) >= MIN_CHUNK_LENGTH:
                    final.append(sub.strip())
            else:
                final.append(chunk)
        return final

    def add_knowledge(self, text, source="manual", source_type="text", metadata=None):
        if not self._ensure_db():
            return []
        chunks = self._chunk_text(text)
        added = []
        conn = get_conn()
        try:
            with conn.cursor() as c:
                for i, chunk in enumerate(chunks):
                    cid = hashlib.md5(f"{chunk}_{source}_{i}".encode()).hexdigest()[:16]
                    c.execute(
                        """INSERT INTO knowledge (id,content,source,source_type,chunk_index,metadata)
                           VALUES (%s,%s,%s,%s,%s,%s) ON CONFLICT (id) DO NOTHING RETURNING id""",
                        (cid, chunk, source, source_type, i, json.dumps(metadata or {}))
                    )
                    if c.fetchone():
                        added.append({"id": cid, "preview": chunk[:100] + "..."})
                        self.knowledge_ids.append(cid)
                        self.knowledge_texts.append(chunk)
            conn.commit()
        except Exception as e:
            print(f"⚠️  Error add_knowledge: {e}")
            conn.rollback()
        finally:
            put_conn(conn)
        if added:
            self._fit_models()
        return added

    def query(self, question, top_k=TOP_K_RESULTS):
        if not self._ensure_db():
            return {"error": "Base de datos no disponible", "answer_chunks": [], "confidence": 0}

        if not self.is_fitted or len(self.knowledge_texts) < 2:
            return {"answer_chunks": [], "confidence": 0.0,
                    "message": "Aún no tengo suficiente conocimiento. Aliméntame con documentos primero.",
                    "synthesized_answer": "Aún no tengo suficiente conocimiento. Ve a la pestaña 'Enseñar' y aliméntame con texto, URLs o documentos."}

        q_tfidf = self.vectorizer.transform([question])
        q_tensor = torch.FloatTensor(q_tfidf.toarray())
        self.encoder.eval()
        with torch.no_grad():
            _, q_emb = self.encoder(q_tensor)

        sims = cosine_similarity(q_emb.numpy(), self.knowledge_embeddings.numpy())[0]

        if self.ranker:
            self.ranker.eval()
            with torch.no_grad():
                attn_scores, _, _ = self.ranker(q_emb.squeeze(0), self.knowledge_embeddings)
            combined = 0.7 * sims + 0.3 * attn_scores.numpy()
        else:
            combined = sims

        top_idx = np.argsort(combined)[-top_k:][::-1]
        results = []
        conn = get_conn()
        try:
            with conn.cursor() as c:
                for idx in top_idx:
                    kid = self.knowledge_ids[idx]
                    c.execute("UPDATE knowledge SET access_count=access_count+1 WHERE id=%s", (kid,))
                    c.execute("SELECT content,source,source_type FROM knowledge WHERE id=%s", (kid,))
                    row = c.fetchone()
                    if row:
                        results.append({"id": kid, "content": row[0], "source": row[1],
                                        "source_type": row[2], "relevance_score": round(float(combined[idx]), 4)})
            conn.commit()
        finally:
            put_conn(conn)

        iid = self._log_interaction(question, [r["id"] for r in results])
        avg = np.mean([r["relevance_score"] for r in results]) if results else 0

        return {"interaction_id": iid, "answer_chunks": results, "confidence": round(float(avg), 4),
                "total_knowledge": len(self.knowledge_texts),
                "synthesized_answer": self._synthesize(results)}

    def _synthesize(self, chunks):
        if not chunks:
            return "No tengo información suficiente para responder."
        parts, seen = [], set()
        for ch in sorted(chunks, key=lambda x: x["relevance_score"], reverse=True):
            h = hashlib.md5(ch["content"][:50].encode()).hexdigest()
            if h not in seen and ch["relevance_score"] > 0.1:
                seen.add(h)
                parts.append(ch["content"])
        return "\n\n".join(parts[:3]) if parts else "Confianza insuficiente. Reformula la pregunta."

    def _log_interaction(self, question, chunk_ids):
        conn = get_conn()
        iid = None
        try:
            with conn.cursor() as c:
                c.execute("INSERT INTO interactions (question,answer_chunks) VALUES (%s,%s) RETURNING id",
                          (question, json.dumps(chunk_ids)))
                iid = c.fetchone()[0]
            conn.commit()
        except Exception as e:
            print(f"⚠️  Error log: {e}")
            conn.rollback()
        finally:
            put_conn(conn)
        return iid

    def get_stats(self):
        if not self._ensure_db():
            return {"error": "BD no disponible"}
        conn = get_conn()
        try:
            with conn.cursor() as c:
                c.execute("SELECT COUNT(*) FROM knowledge")
                total = c.fetchone()[0]
                c.execute("SELECT COUNT(DISTINCT source) FROM knowledge")
                sources = c.fetchone()[0]
                c.execute("SELECT COUNT(*) FROM interactions")
                interactions = c.fetchone()[0]
                c.execute("SELECT source_type, COUNT(*) FROM knowledge GROUP BY source_type")
                by_type = dict(c.fetchall())
                c.execute("SELECT source,COUNT(*) as n FROM knowledge GROUP BY source ORDER BY n DESC LIMIT 10")
                top = [{"source": r[0], "chunks": r[1]} for r in c.fetchall()]
                c.execute("SELECT COUNT(*), AVG(loss) FROM training_log")
                ti = c.fetchone()
            return {"total_chunks": total, "total_sources": sources, "total_interactions": interactions,
                    "knowledge_by_type": by_type, "top_sources": top,
                    "training_sessions": ti[0] or 0,
                    "avg_loss": round(ti[1], 6) if ti[1] else None,
                    "model_status": "trained" if self.is_fitted else "untrained",
                    "database": "PostgreSQL"}
        finally:
            put_conn(conn)

    def provide_feedback(self, interaction_id, score):
        if not self._ensure_db():
            return {"error": "BD no disponible"}
        conn = get_conn()
        try:
            with conn.cursor() as c:
                c.execute("UPDATE interactions SET feedback=%s WHERE id=%s", (score, interaction_id))
                c.execute("SELECT question,answer_chunks FROM interactions WHERE id=%s", (interaction_id,))
                row = c.fetchone()
            conn.commit()
        finally:
            put_conn(conn)
        if row and self.is_fitted:
            question, cids = row
            if isinstance(cids, str):
                cids = json.loads(cids)
            q_t = torch.FloatTensor(self.vectorizer.transform([question]).toarray())
            self.encoder.eval()
            with torch.no_grad():
                _, q_emb = self.encoder(q_t)
            self._train_from_feedback(q_emb.squeeze(0), cids, score)
        return {"status": "ok", "message": f"Feedback registrado: {score}"}

    def save_models(self):
        if not self.encoder or not self.ranker:
            return False
        buf = io.BytesIO()
        torch.save({"encoder": self.encoder.state_dict(), "ranker": self.ranker.state_dict()}, buf)
        conn = get_conn()
        try:
            with conn.cursor() as c:
                c.execute("""INSERT INTO model_weights (id,weights,updated_at) VALUES ('main',%s,NOW())
                             ON CONFLICT (id) DO UPDATE SET weights=%s, updated_at=NOW()""",
                          (psycopg2.Binary(buf.getvalue()), psycopg2.Binary(buf.getvalue())))
            conn.commit()
            return True
        finally:
            put_conn(conn)

    def load_models_from_db(self):
        if not self.encoder or not self.ranker:
            return False
        conn = get_conn()
        try:
            with conn.cursor() as c:
                c.execute("SELECT weights FROM model_weights WHERE id='main'")
                row = c.fetchone()
            if row:
                buf = io.BytesIO(bytes(row[0]))
                cp = torch.load(buf, map_location='cpu', weights_only=True)
                self.encoder.load_state_dict(cp["encoder"])
                self.ranker.load_state_dict(cp["ranker"])
                return True
            return False
        finally:
            put_conn(conn)


# =============================================================================
# INSTANCIA GLOBAL (lazy init - no crash si BD no disponible al arrancar)
# =============================================================================
engine = KnowledgeEngine()


# =============================================================================
# API ROUTES
# =============================================================================

@app.route("/")
def home():
    return jsonify({
        "name": "NeuroBot API",
        "version": "2.1.0",
        "database": "PostgreSQL",
        "endpoints": {
            "GET /health": "Health check",
            "POST /learn/text": "Aprender de texto",
            "POST /learn/document": "Aprender de PDF/DOCX/TXT",
            "POST /learn/url": "Aprender de web",
            "POST /ask": "Preguntar",
            "POST /feedback": "Dar feedback (-1 a 1)",
            "GET /stats": "Estadísticas",
            "GET /knowledge": "Listar conocimiento",
            "POST /save": "Guardar modelo en BD",
            "POST /load": "Cargar modelo de BD"
        }
    })


@app.route("/health")
def health():
    try:
        conn = get_conn()
        with conn.cursor() as c:
            c.execute("SELECT 1")
        put_conn(conn)
        # Init tables on first health check
        engine.init_db()
        return jsonify({"status": "healthy", "db": "connected", "model": engine.is_fitted})
    except Exception as e:
        return jsonify({"status": "unhealthy", "error": str(e)}), 500


@app.route("/learn/text", methods=["POST"])
def learn_text():
    data = request.get_json()
    if not data or "text" not in data:
        return jsonify({"error": "Campo 'text' requerido"}), 400
    added = engine.add_knowledge(data["text"], source=data.get("source", "input_directo"),
                                  source_type="text", metadata=data.get("metadata", {}))
    return jsonify({"status": "learned", "chunks_added": len(added), "chunks": added,
                    "total_knowledge": len(engine.knowledge_texts)})


@app.route("/learn/document", methods=["POST"])
def learn_document():
    if "file" not in request.files:
        return jsonify({"error": "Campo 'file' requerido"}), 400
    file = request.files["file"]
    fname = file.filename.lower()
    tmp = f"/tmp/{file.filename}"
    file.save(tmp)
    try:
        if fname.endswith(".pdf"):
            text, stype = extract_text_from_pdf(tmp), "pdf"
        elif fname.endswith((".docx", ".doc")):
            text, stype = extract_text_from_docx(tmp), "docx"
        elif fname.endswith(".txt"):
            with open(tmp, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            stype = "txt"
        else:
            return jsonify({"error": f"Formato no soportado: {fname}"}), 400
        added = engine.add_knowledge(text, source=file.filename, source_type=stype,
                                      metadata={"original_filename": file.filename})
        return jsonify({"status": "learned", "source": file.filename, "source_type": stype,
                        "text_length": len(text), "chunks_added": len(added), "chunks": added,
                        "total_knowledge": len(engine.knowledge_texts)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if os.path.exists(tmp):
            os.remove(tmp)


@app.route("/learn/url", methods=["POST"])
def learn_url():
    data = request.get_json()
    if not data or "url" not in data:
        return jsonify({"error": "Campo 'url' requerido"}), 400
    try:
        text = extract_text_from_url(data["url"])
        added = engine.add_knowledge(text, source=data["url"], source_type="web", metadata={"url": data["url"]})
        return jsonify({"status": "learned", "source": data["url"], "text_length": len(text),
                        "chunks_added": len(added), "chunks": added,
                        "total_knowledge": len(engine.knowledge_texts)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/ask", methods=["POST"])
def ask():
    data = request.get_json()
    if not data or "question" not in data:
        return jsonify({"error": "Campo 'question' requerido"}), 400
    q = data["question"]
    if data.get("learn_from_question", True):
        engine.add_knowledge(f"Pregunta frecuente: {q}", source="user_questions", source_type="interaction")
    return jsonify(engine.query(q, top_k=data.get("top_k", TOP_K_RESULTS)))


@app.route("/feedback", methods=["POST"])
def feedback():
    data = request.get_json()
    if not data or "interaction_id" not in data or "score" not in data:
        return jsonify({"error": "Campos 'interaction_id' y 'score' requeridos"}), 400
    return jsonify(engine.provide_feedback(data["interaction_id"], max(-1, min(1, float(data["score"])))))


@app.route("/stats")
def stats():
    return jsonify(engine.get_stats())


@app.route("/knowledge")
def list_knowledge():
    if not engine._ensure_db():
        return jsonify({"error": "BD no disponible"}), 500
    page = request.args.get("page", 1, type=int)
    pp = request.args.get("per_page", 20, type=int)
    st = request.args.get("type", None)
    conn = get_conn()
    try:
        with conn.cursor() as c:
            if st:
                c.execute("SELECT id,content,source,source_type,created_at,access_count FROM knowledge WHERE source_type=%s ORDER BY created_at DESC LIMIT %s OFFSET %s", (st, pp, (page-1)*pp))
            else:
                c.execute("SELECT id,content,source,source_type,created_at,access_count FROM knowledge ORDER BY created_at DESC LIMIT %s OFFSET %s", (pp, (page-1)*pp))
            rows = c.fetchall()
            c.execute("SELECT COUNT(*) FROM knowledge")
            total = c.fetchone()[0]
        return jsonify({"page": page, "per_page": pp, "total": total,
            "knowledge": [{"id": r[0], "preview": r[1][:200], "source": r[2], "source_type": r[3],
                           "created_at": r[4].isoformat() if r[4] else None, "access_count": r[5]} for r in rows]})
    finally:
        put_conn(conn)


@app.route("/save", methods=["POST"])
def save_models():
    if engine.save_models():
        return jsonify({"status": "saved", "storage": "PostgreSQL"})
    return jsonify({"error": "No hay modelos para guardar"}), 400


@app.route("/load", methods=["POST"])
def load_models():
    if engine.load_models_from_db():
        return jsonify({"status": "loaded", "storage": "PostgreSQL"})
    return jsonify({"error": "No hay pesos guardados"}), 404


# =============================================================================
# RUN
# =============================================================================
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    print(f"\n  NeuroBot API v2.1 | PostgreSQL | Puerto: {port}\n")
    app.run(host="0.0.0.0", port=port, debug=os.environ.get("FLASK_DEBUG", "false").lower() == "true")
